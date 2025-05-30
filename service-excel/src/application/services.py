import os
import io
import tempfile
import asyncio
import uuid
import json
import pandas as pd
import zipfile
import xlsxwriter
import openpyxl
import pandas as pd
import matplotlib.pyplot as plt
from matplotlib.backends.backend_pdf import PdfPages
from typing import List, Dict, Any, Optional, Tuple
from datetime import datetime
from fastapi import UploadFile, BackgroundTasks
import logging
import hashlib

from application.dto import CreateDocumentDTO, CreateTemplateDTO, TemplateDataDTO, MergeDocumentsDTO, BatchProcessingDTO
from domain.models import ExcelDocumentInfo, ExcelTemplateInfo, BatchProcessingInfo, MergeInfo, ExcelDocumentCreate, ExcelDocumentUpdate
from domain.exceptions import DocumentNotFoundException, TemplateNotFoundException, StorageException
from domain.exceptions import ConversionException, TemplateApplicationException, MergeException
from infrastructure.repository import ExcelDocumentRepository, ExcelTemplateRepository, BatchProcessingRepository, \
    MergeRepository
from infrastructure.minio_client import MinioClient
from infrastructure.rabbitmq_client import RabbitMQClient
from core.config import settings

from openpyxl import load_workbook, Workbook
from openpyxl.utils import get_column_letter
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.table import WD_TABLE_ALIGNMENT

logger = logging.getLogger(__name__)

TEMP_FILE_DIR = "temp/excel_uploads"
os.makedirs(TEMP_FILE_DIR, exist_ok=True)

EXCEL_BUCKET_NAME = "excel-documents"

def _calculate_checksum(file_path: str, hash_algo: str = 'sha256') -> str:
    hasher = hashlib.new(hash_algo)
    with open(file_path, 'rb') as f:
        while chunk := f.read(8192):
            hasher.update(chunk)
    return hasher.hexdigest()

async def _cleanup_temp_file(file_path: str):
    try:
        if os.path.exists(file_path):
            os.remove(file_path)
            logger.info(f"Cleaned up temp file: {file_path}")
    except Exception as e:
        logger.error(f"Error cleaning up temp file {file_path}: {e}")

class ExcelDocumentService:
    """
    Service xử lý tài liệu Excel.
    """

    def __init__(
            self,
            document_repository: ExcelDocumentRepository,
            minio_client: MinioClient,
            rabbitmq_client: Optional[RabbitMQClient] = None
    ):
        """
        Khởi tạo service.

        Args:
            document_repository: Repository để làm việc với tài liệu Excel.
            minio_client: Client MinIO để lưu trữ tài liệu.
            rabbitmq_client: Client RabbitMQ để gửi tin nhắn (optional).
        """
        self.repository = document_repository
        self.minio_client = minio_client
        self.rabbitmq_client = rabbitmq_client

    async def _extract_excel_metadata(self, file_path: str) -> Dict[str, Any]:
        doc_metadata = {}
        try:
            workbook = openpyxl.load_workbook(file_path, read_only=True, data_only=False)
            doc_metadata['sheet_count'] = len(workbook.sheetnames)
            doc_metadata['sheet_names'] = workbook.sheetnames
            if workbook.properties:
                doc_metadata['title_from_properties'] = workbook.properties.title
                doc_metadata['creator'] = workbook.properties.creator
                doc_metadata['last_modified_by'] = workbook.properties.lastModifiedBy
            workbook.close()
            logger.info(f"Extracted metadata from {file_path}: {doc_metadata}")
        except Exception as e:
            logger.warning(f"Could not extract metadata from Excel file {file_path}: {e}", exc_info=True)
        return doc_metadata

    async def convert_to_pdf(self, user_id: str,
                             doc_id: str,
                             background_tasks: BackgroundTasks
                            ) -> Tuple[str, str]:
        """
        Chuyển đổi tài liệu Excel sang PDF.
        Sử dụng pandas để đọc từng sheet và matplotlib để vẽ thành PDF.
        """
        logger.info(f"User {user_id} requesting PDF conversion for doc_id={doc_id}.")
        
        doc_info = await self.get_document_by_id(doc_id, user_id)
        if not doc_info.storage_path or doc_info.file_type not in [
            "application/vnd.ms-excel", 
            "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
        ]:
            raise ConversionException(f"Document {doc_id} is not a valid Excel file or has no storage path.")

        temp_excel_path, _, _ = await self.download_document_content(doc_id, user_id, background_tasks)

        pdf_filename = os.path.splitext(doc_info.original_filename or doc_id)[0] + ".pdf"
        
        safe_pdf_filename = "".join(c if c.isalnum() or c in ('.', '-', '_') else '_' for c in pdf_filename)
        temp_pdf_fd, temp_pdf_path = tempfile.mkstemp(suffix=f"_{safe_pdf_filename}", dir=TEMP_FILE_DIR)
        os.close(temp_pdf_fd)
        
        try:
            xls = pd.ExcelFile(temp_excel_path)
            with PdfPages(temp_pdf_path) as pdf_pages:
                for sheet_name in xls.sheet_names:
                    try:
                        df = pd.read_excel(xls, sheet_name=sheet_name)
                        if df.empty:
                            logger.info(f"Sheet '{sheet_name}' in {doc_id} is empty, skipping in PDF.")
                            fig, ax = plt.subplots(figsize=(11, 8))
                            ax.text(0.5, 0.5, f"Sheet: {sheet_name}\\n(No data)", 
                                    horizontalalignment='center', verticalalignment='center', 
                                    fontsize=12, transform=ax.transAxes)
                            ax.axis('off')
                            pdf_pages.savefig(fig, bbox_inches='tight')
                            plt.close(fig)
                            continue

                        fig, ax = plt.subplots(figsize=(df.shape[1] * 1.5, df.shape[0] * 0.5 + 1))
                        ax.axis('tight')
                        ax.axis('off')
                        
                        the_table = ax.table(cellText=df.values, colLabels=df.columns, loc='center', cellLoc='left')
                        the_table.auto_set_font_size(False)
                        the_table.set_fontsize(8)
                        the_table.scale(1, 1.5)

                        plt.title(sheet_name, fontsize=12)
                        pdf_pages.savefig(fig, bbox_inches='tight')
                        plt.close(fig)
                    except Exception as e_sheet:
                        logger.error(f"Error processing sheet '{sheet_name}' for PDF conversion of {doc_id}: {e_sheet}", exc_info=True)
                        fig, ax = plt.subplots(figsize=(11,8))
                        ax.text(0.5, 0.5, f"Error processing sheet: {sheet_name}\\n{str(e_sheet)[:100]}",
                                color='red', horizontalalignment='center', verticalalignment='center',
                                fontsize=10, transform=ax.transAxes)
                        ax.axis('off')
                        pdf_pages.savefig(fig, bbox_inches='tight')
                        plt.close(fig)

            background_tasks.add_task(_cleanup_temp_file, temp_pdf_path)
            logger.info(f"Successfully converted {doc_id} to PDF at {temp_pdf_path} for user {user_id}.")
            return temp_pdf_path, pdf_filename

        except Exception as e:
            await _cleanup_temp_file(temp_pdf_path)
            logger.error(f"Failed to convert Excel {doc_id} to PDF for user {user_id}: {e}", exc_info=True)
            raise ConversionException(f"Could not convert Excel to PDF: {e}")

    async def convert_to_word(self, user_id: str, doc_id: str, background_tasks: BackgroundTasks) -> ExcelDocumentInfo:
        """
        Chuyển đổi tài liệu Excel sang Word.
        Lấy file từ MinIO, xử lý, lưu file Word mới vào MinIO và DB.
        """
        logger.info(f"User {user_id} requesting Word conversion for doc_id={doc_id}.")
        
        excel_doc_info = await self.get_document_by_id(doc_id, user_id)
        if not excel_doc_info.storage_path or excel_doc_info.file_type not in [
            "application/vnd.ms-excel", 
            "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
        ]:
            raise ConversionException(f"Document {doc_id} is not a valid Excel file or has no storage path.")

        temp_excel_path, _, _ = await self.download_document_content(doc_id, user_id, background_tasks)

        word_original_filename = os.path.splitext(excel_doc_info.original_filename or doc_id)[0] + ".docx"
        
        safe_word_filename = "".join(c if c.isalnum() or c in ('.', '-', '_') else '_' for c in word_original_filename)
        temp_word_fd, temp_word_path = tempfile.mkstemp(suffix=f"_{safe_word_filename}", dir=TEMP_FILE_DIR)
        os.close(temp_word_fd)

        try:
            wb = load_workbook(temp_excel_path)
            doc = Document()
            doc.add_heading(os.path.splitext(excel_doc_info.original_filename or "Converted Document")[0], 0)

            for sheet_name in wb.sheetnames:
                sheet = wb[sheet_name]
                doc.add_heading(f"Sheet: {sheet_name}", level=2)
                if sheet.max_row == 0 or sheet.max_column == 0:
                    doc.add_paragraph("(Sheet is empty)")
                    continue

                table = doc.add_table(rows=1, cols=sheet.max_column)
                table.style = 'Table Grid'
                hdr_cells = table.rows[0].cells
                for col_idx, cell in enumerate(sheet[1]):
                    hdr_cells[col_idx].text = str(cell.value if cell.value is not None else '')

                for i, row in enumerate(sheet.iter_rows(min_row=2)):
                    row_cells = table.add_row().cells
                    for j, cell in enumerate(row):
                        value = cell.value if cell.value is not None else ''
                        row_cells[j].text = str(value)

                doc.save(temp_word_path)
            word_file_size = os.path.getsize(temp_word_path)
            word_checksum = _calculate_checksum(temp_word_path)
            
            background_tasks.add_task(_cleanup_temp_file, temp_word_path)

            word_storage_id = str(uuid.uuid4())
            word_object_name = f"{user_id}/{word_storage_id}/{safe_word_filename}"
            
            await self.minio_client.upload_file(
                file_path=temp_word_path,
                bucket_name=EXCEL_BUCKET_NAME,
                object_name=word_object_name,
                content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )

            
            import json
            
            generic_word_info = {
                "id": str(uuid.uuid4()),
                "storage_id": word_storage_id,
                "document_category": "word",
                "title": os.path.splitext(word_original_filename)[0],
                "description": f"Word document converted from Excel: {excel_doc_info.title or doc_id}",
                "original_filename": word_original_filename,
                "file_type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                "file_size": word_file_size,
                "doc_metadata": {
                    "converted_from": word_original_filename,
                    "original_excel_filename": word_original_filename,
                    "conversion_method": "openpyxl_to_docx"
                },
                "user_id": user_id,
                "created_at": datetime.utcnow(),
                "updated_at": datetime.utcnow(),
                "version": 1,
                "checksum": word_checksum
            }
            
            async with self.repository.pool.acquire() as connection:
                query = """
                    INSERT INTO documents (
                        id, storage_id, document_category, title, description,
                        file_size, file_type, storage_path, original_filename, 
                        doc_metadata, created_at, updated_at, user_id, version, checksum
                    ) VALUES ($1, $2, $3, $4, $5, $6, $7, $8, $9, $10, $11, $12, $13, $14, $15)
                    RETURNING id;
                """
                saved_word_doc_id = await connection.fetchval(
                    query,
                    generic_word_info["id"],
                    generic_word_info["storage_id"],
                    generic_word_info["document_category"],
                    generic_word_info["title"],
                    generic_word_info["description"],
                    generic_word_info["file_size"],
                    generic_word_info["file_type"],
                    word_object_name,
                    json.dumps(generic_word_info["doc_metadata"]),
                    generic_word_info["created_at"],
                    generic_word_info["updated_at"],
                    generic_word_info["user_id"],
                    generic_word_info["version"],
                    generic_word_info["checksum"]
                )
            
            class WordDocumentResult:
                def __init__(self, word_info):
                    self.id = word_info["id"]
                    self.storage_id = word_info["storage_id"]
                    self.document_category = word_info["document_category"]
                    self.title = word_info["title"]
                    self.description = word_info["description"]
                    self.original_filename = word_info["original_filename"]
                    self.file_size = word_info["file_size"]
                    self.file_type = word_info["file_type"]
                    self.doc_metadata = word_info["doc_metadata"]
                    self.user_id = word_info["user_id"]
                    self.checksum = word_info["checksum"]
                    self.version = word_info["version"]
                    self.created_at = word_info["created_at"]
                    self.updated_at = word_info["updated_at"]
            
            saved_word_doc = WordDocumentResult(generic_word_info)
            logger.info(f"Successfully converted Excel {doc_id} to Word {saved_word_doc.id} for user {user_id}.")
            return saved_word_doc

        except Exception as e:
            if os.path.exists(temp_word_path):
                await _cleanup_temp_file(temp_word_path)
            logger.error(f"Failed to convert Excel {doc_id} to Word for user {user_id}: {e}", exc_info=True)
            raise ConversionException(f"Could not convert Excel to Word: {e}")

    async def merge_documents(self, user_id: str, dto: MergeDocumentsDTO, background_tasks: BackgroundTasks) -> ExcelDocumentInfo:
        """
        Gộp nhiều tài liệu Excel thành một.
        Tất cả document_ids phải thuộc về user_id.
        File kết quả được lưu vào MinIO và doc_metadata vào DB.
        """
        logger.info(f"User {user_id} requesting to merge documents: {dto.document_ids} into {dto.output_filename}")
        if not dto.document_ids:
            raise ValueError("No document IDs provided for merging.")
        if not dto.output_filename:
            raise ValueError("Output filename must be provided for merging.")
        if not dto.output_filename.lower().endswith(('.xlsx', '.xls')):
            dto.output_filename += ".xlsx"

        temp_file_paths_to_cleanup = []
        
        try:
            source_documents_info = []
            for doc_id_to_merge in dto.document_ids:
                doc_info = await self.get_document_by_id(doc_id_to_merge, user_id)
                if not doc_info.storage_path or doc_info.file_type not in [
                     "application/vnd.ms-excel", 
                     "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
                ]:
                    raise ConversionException(f"Document {doc_id_to_merge} is not a valid Excel file for merging.")
                source_documents_info.append(doc_info)

            merged_workbook = Workbook()
            if "Sheet" in merged_workbook.sheetnames and len(merged_workbook.sheetnames) == 1:
                merged_workbook.remove(merged_workbook.active)

            for doc_info in source_documents_info:
                temp_excel_path, _, _ = await self.download_document_content(doc_info.id, user_id, background_tasks)

                src_wb = load_workbook(temp_excel_path)
                for sheet_name in src_wb.sheetnames:
                    src_sheet = src_wb[sheet_name]

                    new_sheet_name_base = f"{os.path.splitext(doc_info.original_filename or doc_info.id)[0]}_{sheet_name}"
                    new_sheet_name = new_sheet_name_base[:31]
                    idx = 1
                    while new_sheet_name in merged_workbook.sheetnames:
                        suffix = f"_{idx}"
                        new_sheet_name = new_sheet_name_base[:31-len(suffix)] + suffix
                        idx += 1
                        if len(new_sheet_name) > 31: 
                            new_sheet_name = new_sheet_name[:31]

                        dest_sheet = merged_workbook.create_sheet(title=new_sheet_name)

                        for row in src_sheet.iter_rows():
                            for cell in row:
                                dest_cell = dest_sheet.cell(row=cell.row, column=cell.column, value=cell.value)
                                if cell.has_style:
                                    dest_cell.font = cell.font.copy()
                                    dest_cell.border = cell.border.copy()
                                    dest_cell.fill = cell.fill.copy()
                                    dest_cell.number_format = cell.number_format
                                    dest_cell.alignment = cell.alignment.copy()
                    
                    for col_letter, dim in src_sheet.column_dimensions.items():
                        dest_sheet.column_dimensions[col_letter].width = dim.width
                        if dim.hidden: 
                            dest_sheet.column_dimensions[col_letter].hidden = True

                    for row_idx, dim in src_sheet.row_dimensions.items():
                        dest_sheet.row_dimensions[row_idx].height = dim.height
                        if dim.hidden: 
                            dest_sheet.row_dimensions[row_idx].hidden = True
                    src_wb.close()
            
            safe_output_filename = "".join(c if c.isalnum() or c in ('.', '-', '_') else '_' for c in dto.output_filename)
            temp_merged_fd, temp_merged_path = tempfile.mkstemp(suffix=f"_{safe_output_filename}", dir=TEMP_FILE_DIR)
            os.close(temp_merged_fd)
            temp_file_paths_to_cleanup.append(temp_merged_path)

            merged_workbook.save(temp_merged_path)
            merged_file_size = os.path.getsize(temp_merged_path)
            merged_checksum = _calculate_checksum(temp_merged_path)
            
            background_tasks.add_task(_cleanup_temp_file, temp_merged_path)

            merged_storage_id = str(uuid.uuid4())
            merged_object_name = f"{user_id}/{merged_storage_id}/{safe_output_filename}"
            
            await self.minio_client.upload_file(
                file_path=temp_merged_path,
                bucket_name=EXCEL_BUCKET_NAME,
                object_name=merged_object_name,
                content_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
            )

            merged_doc_db_info = ExcelDocumentInfo(
                id=str(uuid.uuid4()),
                storage_id=merged_storage_id,
                document_category="excel",
                title=os.path.splitext(dto.output_filename)[0],
                description=f"Merged Excel document from {len(source_documents_info)} files.",
                file_size=merged_file_size,
                file_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                storage_path=merged_object_name,
                original_filename=dto.output_filename,
                doc_metadata={
                    "merged_from_document_ids": [doc.id for doc in source_documents_info],
                    "merged_from_document_titles": [doc.title or doc.original_filename for doc in source_documents_info],
                    "merged_sheet_names": merged_workbook.sheetnames
                },
                user_id=user_id,
                checksum=merged_checksum,
                sheet_count=len(merged_workbook.sheetnames),
                version=1,
                created_at=datetime.utcnow(),
                updated_at=datetime.utcnow()
            )
            saved_merged_doc = await self.repository.save(merged_doc_db_info)
            logger.info(f"Successfully merged documents into {saved_merged_doc.id} for user {user_id}.")
            return saved_merged_doc

        except Exception as e:
            logger.error(f"Failed to merge documents for user {user_id}: {e}", exc_info=True)
            for temp_path in temp_file_paths_to_cleanup:
                if os.path.exists(temp_path): 
                    await _cleanup_temp_file(temp_path)
            if isinstance(e, (DocumentNotFoundException, ConversionException, StorageException)):
                raise
            raise MergeException(f"Could not merge documents: {e}")

    async def merge_documents_async(self, user_id: str, task_id: str, dto: MergeDocumentsDTO) -> None:
        """
        Gộp nhiều tài liệu Excel thành một (bất đồng bộ).
        Args:
            user_id: ID của người dùng.
            task_id: ID của tác vụ.
            dto: DTO chứa thông tin gộp tài liệu.
        """
        logger.info(f"User {user_id} submitting async merge task {task_id} for docs {dto.document_ids}")
        if not self.rabbitmq_client:
            logger.error("RabbitMQ client not configured for async merge.")
            raise ConnectionError("RabbitMQ client not configured for async merge.")
        
        if not dto.document_ids:
            raise ValueError("No document IDs provided for async merging.")
        if not dto.output_filename:
            raise ValueError("Output filename must be provided for async merging.")

        if dto.document_ids:
            try:
                await self.get_document_by_id(dto.document_ids[0], user_id)
            except DocumentNotFoundException:
                logger.error(f"First document {dto.document_ids[0]} for async merge not found or not accessible by user {user_id}")
                raise DocumentNotFoundException(f"Initial check failed: Document {dto.document_ids[0]} not accessible.")

        message_body = {
            "task_id": task_id,
            "user_id": user_id,
            "document_ids": dto.document_ids,
            "output_filename": dto.output_filename,
            "task_type": "merge_excel_documents"
        }
        
        try:
            await self.rabbitmq_client.publish_message( 
                exchange_name=settings.RABBITMQ_EXCHANGE_TASKS or 'tasks_exchange', # Use settings if available
                routing_key=settings.RABBITMQ_ROUTING_KEY_EXCEL_MERGE or 'excel.tasks.merge', # Use settings
                message_body=json.dumps(message_body)
            )
            logger.info(f"Async merge task {task_id} for user {user_id} published successfully.")
        except Exception as e:
            logger.error(f"Failed to publish async merge task {task_id} for user {user_id}: {e}", exc_info=True)
            raise ConnectionError(f"Failed to publish async merge task: {e}")

    async def get_merge_status(self, user_id: str, task_id: str) -> Dict[str, Any]:
        """
        Lấy trạng thái gộp tài liệu.
        Args:
            user_id: ID của người dùng (để kiểm tra quyền xem task).
            task_id: ID của tác vụ gộp tài liệu.
        Returns:
            Dict chứa thông tin trạng thái.
        """
        logger.info(f"User {user_id} fetching status for merge task {task_id}")
        raise NotImplementedError(
            "Task status retrieval is not implemented. "
            "A system for tracking asynchronous task statuses is required. "
            "The worker processing the RabbitMQ message should update this status."
        )

    async def process_document_async(self, user_id: str, document_id: str, task_type: str = "convert_to_pdf") -> Dict[str, Any]:
        """
        Xử lý tài liệu bất đồng bộ (ví dụ: convert to PDF).
        Args:
            user_id: ID của người dùng.
            document_id: ID của tài liệu.
            task_type: Loại tác vụ cần thực hiện (e.g., "convert_to_pdf", "convert_to_word").
        Returns:
            A dict containing the task_id submitted.
        """
        logger.info(f"User {user_id} requesting async processing task '{task_type}' for document {document_id}")
        
        await self.get_document_by_id(document_id, user_id)
        
        if not self.rabbitmq_client:
            logger.error(f"RabbitMQ client not configured for async document processing task '{task_type}'.")
            raise ConnectionError("RabbitMQ client not configured.")

        task_id = str(uuid.uuid4())
        message_body = {
                "task_id": task_id,
            "user_id": user_id,
            "document_id": document_id,
            "task_type": task_type
        }
        
        routing_key_map = {
            "convert_to_pdf": getattr(settings, 'RABBITMQ_ROUTING_KEY_EXCEL_TO_PDF', 'excel.tasks.convert.pdf'),
            "convert_to_word": getattr(settings, 'RABBITMQ_ROUTING_KEY_EXCEL_TO_WORD', 'excel.tasks.convert.word'),
        }
        routing_key = routing_key_map.get(task_type)
        if not routing_key:
            raise ValueError(f"Unsupported task_type '{task_type}' for async processing.")

        try:
            await self.rabbitmq_client.publish_message(
                exchange_name=getattr(settings, 'RABBITMQ_EXCHANGE_TASKS', 'tasks_exchange'),
                routing_key=routing_key,
                message_body=json.dumps(message_body)
            )
            logger.info(f"Async processing task {task_id} ('{task_type}') for doc {document_id} (user {user_id}) published successfully.")
            return {"task_id": task_id, "status": "submitted", "message": f"Task '{task_type}' for document {document_id} submitted."}
        except Exception as e:
            logger.error(f"Failed to publish async task {task_id} ('{task_type}') for doc {document_id}: {e}", exc_info=True)
            raise ConnectionError(f"Failed to publish async task: {e}")

    async def upload_document(
        self, 
        user_id: str, 
        file: UploadFile, 
        doc_create_info: Optional[ExcelDocumentCreate],
        background_tasks: BackgroundTasks
    ) -> ExcelDocumentInfo:
        if not file.filename:
            raise ValueError("Filename cannot be empty.")
        safe_filename = "".join(c if c.isalnum() or c in ('.', '-', '_') else '_' for c in file.filename)
        temp_fd, temp_file_path = tempfile.mkstemp(suffix=f"_{safe_filename}", dir=TEMP_FILE_DIR)
        os.close(temp_fd)
        
        file_size = 0
        try:
            content_to_write = await file.read()
            with open(temp_file_path, "wb") as buffer:
                buffer.write(content_to_write)
            file_size = len(content_to_write)
        except Exception as e:
            logger.error(f"Failed to write temp file {temp_file_path}: {e}", exc_info=True)
            await _cleanup_temp_file(temp_file_path)
            raise StorageException(f"Failed to save temporary file: {e}")
        
        background_tasks.add_task(_cleanup_temp_file, temp_file_path)

        extracted_meta = await self._extract_excel_metadata(temp_file_path)
        checksum = _calculate_checksum(temp_file_path)
        
        storage_id = str(uuid.uuid4())
        object_name = f"{user_id}/{storage_id}/{safe_filename}"

        try:
            returned_storage_key = await self.minio_client.upload_file(
                file_path=temp_file_path, 
                bucket_name=EXCEL_BUCKET_NAME, 
                object_name=object_name,
                content_type=file.content_type
            )
            if returned_storage_key != object_name:
                logger.warning(f"Storage client returned key '{returned_storage_key}' which differs from expected '{object_name}'. Using expected.")
        except Exception as e:
            logger.error(f"Failed to upload {safe_filename} to storage: {e}", exc_info=True)
            raise StorageException(f"Failed to upload to storage: {e}")

        title = (doc_create_info.title if doc_create_info and doc_create_info.title 
                 else extracted_meta.get('title_from_properties') or os.path.splitext(safe_filename)[0])
        description = doc_create_info.description if doc_create_info and doc_create_info.description else None
        
        db_doc_info = ExcelDocumentInfo(
            id=str(uuid.uuid4()),
            storage_id=storage_id,
            document_category="excel",
            title=title,
            description=description,
            file_size=file_size,
            file_type=file.content_type or "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            storage_path=object_name, 
            original_filename=file.filename,
            doc_metadata=extracted_meta,
            user_id=user_id,
            checksum=checksum,
            sheet_count=extracted_meta.get('sheet_count'),
            version=1,
            created_at=datetime.utcnow(),
            updated_at=datetime.utcnow()
        )

        try:
            saved_document = await self.repository.save(db_doc_info)
            logger.info(f"Document {saved_document.id} for user {user_id} saved to DB with storage key {object_name}.")
            return saved_document
        except Exception as e:
            logger.error(f"Failed to save document metadata for {safe_filename} to DB: {e}", exc_info=True)
            logger.info(f"Attempting to delete orphaned file from storage: {EXCEL_BUCKET_NAME}/{object_name}")
            try:
                await self.minio_client.delete_file(EXCEL_BUCKET_NAME, object_name)
            except Exception as del_e:
                logger.error(f"Failed to delete orphaned file {EXCEL_BUCKET_NAME}/{object_name}: {del_e}", exc_info=True)
            raise StorageException(f"Failed to save document doc_metadata: {e}")

    async def get_document_by_id(self, doc_id: str, user_id: str) -> ExcelDocumentInfo:
        logger.debug(f"Fetching document {doc_id} for user {user_id}")
        document = await self.repository.get_by_id(doc_id, user_id)
        if not document:
            raise DocumentNotFoundException(f"Document with id {doc_id} not found for user {user_id}")
        return document

    async def download_document_content(
        self, doc_id: str, user_id: str, background_tasks: BackgroundTasks
    ) -> Tuple[str, str, int]: 
        doc_info = await self.get_document_by_id(doc_id, user_id)
        
        object_key = doc_info.storage_path 
        if not object_key:
            raise StorageException(f"Document {doc_id} has no storage path (object key). Cannot download.")

        safe_original_filename = "".join(c if c.isalnum() or c in ('.', '-', '_') else '_' for c in (doc_info.original_filename or doc_id))
        temp_fd, temp_download_path = tempfile.mkstemp(suffix=f"_{safe_original_filename}", dir=TEMP_FILE_DIR)
        os.close(temp_fd)

        try:
            await self.minio_client.download_file(
                bucket_name=EXCEL_BUCKET_NAME, 
                object_name=object_key, 
                download_path=temp_download_path
            )
            background_tasks.add_task(_cleanup_temp_file, temp_download_path)
            
            file_size_to_return = doc_info.file_size
            if file_size_to_return is None:
                 try:
                     file_size_to_return = os.path.getsize(temp_download_path)
                 except OSError:
                     logger.warning(f"Could not get size of downloaded file {temp_download_path}")
                     file_size_to_return = 0
            return temp_download_path, doc_info.original_filename or doc_id, file_size_to_return
        except FileNotFoundError:
            await _cleanup_temp_file(temp_download_path)
            logger.error(f"File not found in storage: bucket '{EXCEL_BUCKET_NAME}', object '{object_key}'")
            raise DocumentNotFoundException(f"File for document {doc_id} not found in storage.")
        except Exception as e:
            await _cleanup_temp_file(temp_download_path)
            logger.error(f"Error downloading file {doc_id} from storage: {e}", exc_info=True)
            raise StorageException(f"Could not download file {doc_id}: {e}")

    async def list_documents_by_user(
        self, user_id: str, skip: int, limit: int, 
        search_term: Optional[str] = None, sort_by: str = 'created_at', sort_order: str = 'desc'
    ) -> Tuple[List[ExcelDocumentInfo], int]:
        logger.debug(f"Listing documents for user {user_id} with skip={skip}, limit={limit}, search='{search_term}'")
        return await self.repository.list_by_user_id(user_id, skip, limit, search_term, sort_by, sort_order)

    async def update_document_metadata(
        self, doc_id: str, user_id: str, update_data: ExcelDocumentUpdate
    ) -> ExcelDocumentInfo:
        
        update_dict = update_data.model_dump(exclude_unset=True)
        if not update_dict:
            logger.info(f"No data provided to update document {doc_id}")
            return await self.get_document_by_id(doc_id, user_id)
        
        current_doc = await self.repository.get_by_id(doc_id, user_id)
        if not current_doc:
             raise DocumentNotFoundException(f"Document {doc_id} not found for user {user_id} to update.")

        if 'doc_metadata' in update_dict and update_dict['doc_metadata'] is not None:
            if 'sheet_count' in update_dict['doc_metadata']:
                update_dict['sheet_count'] = update_dict['doc_metadata']['sheet_count']
        
        logger.debug(f"Updating document {doc_id} for user {user_id} with data: {update_dict}")
        updated_doc = await self.repository.update_metadata(doc_id, user_id, update_dict)
        if not updated_doc:
             raise DocumentNotFoundException(f"Failed to update document {doc_id} or document not found after update attempt.")
        return updated_doc

    async def delete_document(
        self, doc_id: str, user_id: str, background_tasks: BackgroundTasks
    ) -> bool:
        doc_to_delete = await self.repository.get_by_id(doc_id, user_id)
        if not doc_to_delete:
            logger.warning(f"Document {doc_id} not found for user {user_id} to delete. No action taken.")
            return False

        object_key_to_delete = doc_to_delete.storage_path
        deleted_from_db = await self.repository.delete(doc_id, user_id)

        if deleted_from_db:
            logger.info(f"Document {doc_id} (user {user_id}) deleted from DB.")
            if object_key_to_delete:
                logger.info(f"Scheduling deletion of storage object: {EXCEL_BUCKET_NAME}/{object_key_to_delete}")
                background_tasks.add_task(
                    self.minio_client.delete_file,
                    bucket_name=EXCEL_BUCKET_NAME, 
                    object_name=object_key_to_delete
                )
            else:
                 logger.warning(f"Document {doc_id} had no storage_path, so no file to delete from storage.")
            return True
        
        logger.warning(f"Failed to delete document {doc_id} (user {user_id}) from DB (it might have been deleted by another process or an error occurred).")
        return False

class ExcelTemplateService:
    """
    Service xử lý mẫu tài liệu Excel.
    """

    def __init__(
            self,
            template_repository: ExcelTemplateRepository,
            minio_client: MinioClient,
            rabbitmq_client: RabbitMQClient
    ):
        """
        Khởi tạo service.

        Args:
            template_repository: Repository để làm việc với mẫu tài liệu
            minio_client: Client MinIO để lưu trữ mẫu tài liệu
            rabbitmq_client: Client RabbitMQ để gửi tin nhắn
        """
        self.template_repository = template_repository
        self.minio_client = minio_client
        self.rabbitmq_client = rabbitmq_client
        self.batch_repository = BatchProcessingRepository()
        self.document_repository = ExcelDocumentRepository(minio_client)

    async def create_template(self, dto: CreateTemplateDTO, content: bytes) -> ExcelTemplateInfo:
        """
        Tạo mẫu tài liệu mới.

        Args:
            dto: DTO chứa thông tin mẫu tài liệu
            content: Nội dung mẫu tài liệu

        Returns:
            Thông tin mẫu tài liệu đã tạo
        """
        sheet_names = await self._get_sheet_names(content)

        template_info = ExcelTemplateInfo(
            name=dto.name,
            description=dto.description,
            category=dto.category,
            original_filename=dto.original_filename,
            file_size=len(content),
            storage_path="",  
            data_fields=dto.data_fields,
            doc_metadata=dto.doc_metadata,
            sheet_names=sheet_names
        )

        template_info = await self.template_repository.save(template_info, content)

        return template_info

    async def _get_sheet_names(self, content: bytes) -> List[str]:
        """
        Lấy danh sách tên sheet từ file Excel.

        Args:
            content: Nội dung file Excel

        Returns:
            Danh sách tên sheet
        """
        try:
            with tempfile.NamedTemporaryFile(delete=False, suffix=".xlsx") as temp_file:
                temp_file.write(content)
                temp_file_path = temp_file.name

            try:
                wb = load_workbook(temp_file_path, read_only=True)
                sheet_names = wb.sheetnames
                wb.close()

                os.unlink(temp_file_path)

                return sheet_names
            except Exception as e:
                if os.path.exists(temp_file_path):
                    os.unlink(temp_file_path)
                print(f"Lỗi khi đọc tên sheet: {str(e)}")
                return []
        except Exception as e:
            print(f"Lỗi khi tạo file tạm: {str(e)}")
            return []

    async def get_templates(self, category: Optional[str] = None, skip: int = 0, limit: int = 10) -> List[
        ExcelTemplateInfo]:
        """
        Lấy danh sách mẫu tài liệu.

        Args:
            category: Danh mục để lọc
            skip: Số mẫu tài liệu bỏ qua
            limit: Số mẫu tài liệu tối đa trả về

        Returns:
            Danh sách mẫu tài liệu
        """
        return await self.template_repository.list(category, skip, limit)

    async def get_template(self, template_id: str) -> Tuple[ExcelTemplateInfo, bytes]:
        """
        Lấy thông tin và nội dung mẫu tài liệu.

        Args:
            template_id: ID của mẫu tài liệu

        Returns:
            Tuple chứa thông tin và nội dung mẫu tài liệu
        """
        return await self.template_repository.get(template_id)

    async def delete_template(self, template_id: str) -> None:
        """
        Xóa mẫu tài liệu.

        Args:
            template_id: ID của mẫu tài liệu
        """
        await self.template_repository.delete(template_id)

    async def apply_template(self, dto: TemplateDataDTO) -> Dict[str, Any]:
        """
        Áp dụng mẫu tài liệu với dữ liệu được cung cấp.

        Args:
            dto: DTO chứa thông tin để áp dụng vào mẫu

        Returns:
            Dict chứa thông tin tài liệu đã tạo
        """
        try:
            template_info, template_content = await self.template_repository.get(dto.template_id)

            with tempfile.NamedTemporaryFile(delete=False, suffix=".xlsx") as temp_template:
                temp_template.write(template_content)
                temp_template_path = temp_template.name

            result_filename = f"{template_info.name}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.xlsx"
            temp_result_path = os.path.join(settings.TEMP_DIR, result_filename)

            try:
                wb = load_workbook(temp_template_path)

                for sheet in wb.worksheets:
                    for row in sheet.iter_rows():
                        for cell in row:
                            if cell.value and isinstance(cell.value, str) and "{{" in cell.value and "}}" in cell.value:
                                for key, value in dto.data.items():
                                    placeholder = f"{{{{{key}}}}}"
                                    if placeholder in cell.value:
                                        cell.value = cell.value.replace(placeholder, str(value))

                wb.save(temp_result_path)

                if dto.output_format.lower() == "pdf":
                    pdf_filename = os.path.splitext(result_filename)[0] + ".pdf"
                    temp_pdf_path = os.path.join(settings.TEMP_DIR, pdf_filename)

                    try:

                        xls = pd.ExcelFile(temp_result_path)
                        with PdfPages(temp_pdf_path) as pdf_pages:
                            for sheet_name in xls.sheet_names:
                                try:
                                    df = pd.read_excel(xls, sheet_name=sheet_name)
                                    if df.empty:
                                        fig, ax = plt.subplots(figsize=(11, 8))
                                        ax.text(0.5, 0.5, f"Sheet: {sheet_name}\\n(No data)", 
                                                horizontalalignment='center', verticalalignment='center', 
                                                fontsize=12, transform=ax.transAxes)
                                        ax.axis('off')
                                        pdf_pages.savefig(fig, bbox_inches='tight')
                                        plt.close(fig)
                                        continue

                                    fig, ax = plt.subplots(figsize=(max(df.shape[1] * 1.5, 8), max(df.shape[0] * 0.5 + 1, 6)))
                                    ax.axis('tight')
                                    ax.axis('off')
                                    
                                    the_table = ax.table(cellText=df.values, colLabels=df.columns, loc='center', cellLoc='left')
                                    the_table.auto_set_font_size(False)
                                    the_table.set_fontsize(8)
                                    the_table.scale(1, 1.5)

                                    
                                    plt.title(f"Template: {template_info.name} - Sheet: {sheet_name}", fontsize=12)
                                    pdf_pages.savefig(fig, bbox_inches='tight')
                                    plt.close(fig)
                                except Exception as e_sheet:
                                    logger.error(f"Error processing sheet '{sheet_name}' for PDF template: {e_sheet}", exc_info=True)

                                    fig, ax = plt.subplots(figsize=(11,8))
                                    ax.text(0.5, 0.5, f"Error processing sheet: {sheet_name}\\n{str(e_sheet)[:100]}",
                                            color='red', horizontalalignment='center', verticalalignment='center',
                                            fontsize=10, transform=ax.transAxes)
                                    ax.axis('off')
                                    pdf_pages.savefig(fig, bbox_inches='tight')
                                    plt.close(fig)

                        with open(temp_pdf_path, "rb") as f:
                            result_content = f.read()

                        os.unlink(temp_pdf_path)

                        result_filename = pdf_filename
                    except Exception as e:
                        raise TemplateApplicationException(f"Lỗi khi chuyển đổi sang PDF: {str(e)}")
                else:
                    with open(temp_result_path, "rb") as f:
                        result_content = f.read()

                os.unlink(temp_template_path)
                os.unlink(temp_result_path)

                document_info = ExcelDocumentInfo(
                    title=f"{template_info.name} - {datetime.now().strftime('%Y-%m-%d')}",
                    description=f"Tài liệu được tạo từ mẫu '{template_info.name}'",
                    original_filename=result_filename,
                    file_size=len(result_content),
                    file_type="application/pdf" if dto.output_format.lower() == "pdf" else "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                    storage_path="",  
                    doc_metadata={
                        "template_id": template_info.id,
                        "template_name": template_info.name,
                        "template_data": dto.data
                    },
                    user_id=dto.user_id
                )

                document_info = await self.document_repository.save(document_info, result_content)

                return {
                    "id": document_info.id,
                    "filename": document_info.original_filename,
                    "file_size": document_info.file_size
                }
            except Exception as e:
                if os.path.exists(temp_template_path):
                    os.unlink(temp_template_path)
                if os.path.exists(temp_result_path):
                    os.unlink(temp_result_path)

                raise TemplateApplicationException(f"Lỗi khi áp dụng mẫu: {str(e)}")
        except TemplateNotFoundException:
            raise
        except Exception as e:
            raise TemplateApplicationException(str(e))

    async def process_batch_async(self, task_id: str, template_id: str, content: bytes, filename: str,
                                  output_format: str) -> None:
        """
        Xử lý batch tài liệu.

        Args:
            task_id: ID của tác vụ
            template_id: ID của mẫu tài liệu
            content: Nội dung file dữ liệu (CSV, Excel)
            filename: Tên file dữ liệu
            output_format: Định dạng đầu ra
        """
        try:
            batch_info = BatchProcessingInfo(
                id=task_id,
                template_id=template_id,
                output_format=output_format
            )

            await self.batch_repository.save(batch_info)

            with tempfile.NamedTemporaryFile(delete=False, suffix=f".{filename.split('.')[-1]}") as temp_file:
                temp_file.write(content)
                temp_file_path = temp_file.name

            try:
                if filename.endswith('.csv'):
                    data_list = pd.read_csv(temp_file_path).to_dict('records')
                elif filename.endswith(('.xlsx', '.xls')):
                    data_list = pd.read_excel(temp_file_path).to_dict('records')
                else:
                    raise TemplateApplicationException(f"Định dạng file không được hỗ trợ: {filename}")

                batch_info.total_documents = len(data_list)
                await self.batch_repository.update(batch_info)

                result_documents = []

                for i, data in enumerate(data_list):
                    try:
                        template_data_dto = TemplateDataDTO(
                            template_id=template_id,
                            data=data,
                            output_format=output_format
                        )

                        result = await self.apply_template(template_data_dto)
                        result_documents.append(result)

                        batch_info.processed_documents = i + 1
                        await self.batch_repository.update(batch_info)
                    except Exception as e:
                        print(f"Lỗi khi xử lý bản ghi thứ {i}: {str(e)}")

                if output_format.lower() == "zip":
                    zip_filename = f"batch_{task_id}.zip"
                    temp_zip_path = os.path.join(settings.TEMP_DIR, zip_filename)

                    with zipfile.ZipFile(temp_zip_path, 'w', zipfile.ZIP_DEFLATED) as zipf:
                        for result in result_documents:
                            document_info, content = await self.document_repository.get(result["id"])
                            zipf.writestr(document_info.original_filename, content)

                    with open(temp_zip_path, "rb") as f:
                        zip_content = f.read()

                    zip_document_info = ExcelDocumentInfo(
                        title=f"Batch {task_id}",
                        description=f"File ZIP chứa {len(result_documents)} tài liệu được tạo từ mẫu",
                        original_filename=zip_filename,
                        file_size=len(zip_content),
                        file_type="application/zip",
                        storage_path="",  
                        doc_metadata={
                            "template_id": template_id,
                            "batch_id": task_id,
                            "total_documents": len(result_documents)
                        }
                    )

                    zip_document_info = await self.document_repository.save(zip_document_info, zip_content)

                    batch_info.status = "completed"
                    batch_info.completed_at = datetime.now()
                    batch_info.result_file_id = zip_document_info.id
                    batch_info.result_file_path = zip_document_info.storage_path
                    await self.batch_repository.update(batch_info)

                    os.unlink(temp_zip_path)
                else:
                    batch_info.status = "completed"
                    batch_info.completed_at = datetime.now()
                    await self.batch_repository.update(batch_info)
            except Exception as e:
                if os.path.exists(temp_file_path):
                    os.unlink(temp_file_path)

                batch_info.status = "failed"
                batch_info.error_message = str(e)
                await self.batch_repository.update(batch_info)

                raise TemplateApplicationException(f"Lỗi khi xử lý batch: {str(e)}")

            os.unlink(temp_file_path)
        except Exception as e:
            try:
                batch_info = await self.batch_repository.get(task_id)
                batch_info.status = "failed"
                batch_info.error_message = str(e)
                await self.batch_repository.update(batch_info)
            except:
                pass

            print(f"Lỗi khi xử lý batch {task_id}: {str(e)}")

    async def get_batch_status(self, batch_id: str) -> Dict[str, Any]:
        """
        Lấy trạng thái xử lý batch.

        Args:
            batch_id: ID của batch

        Returns:
            Dict chứa thông tin trạng thái
        """
        try:
            batch_info = await self.batch_repository.get(batch_id)
            status_data = {
                "task_id": batch_info.id,
                "status": batch_info.status,
                "created_at": batch_info.created_at.isoformat(),
                "total_documents": batch_info.total_documents,
                "processed_documents": batch_info.processed_documents,
                "output_format": batch_info.output_format
            }

            if batch_info.status == "completed" and batch_info.result_file_id:
                status_data["result_file_id"] = batch_info.result_file_id
                status_data["download_url"] = f"/documents/download/{batch_info.result_file_id}"
            elif batch_info.status == "failed" and batch_info.error_message:
                status_data["error_message"] = batch_info.error_message

            return status_data
        except Exception as e:
            return {
                "task_id": batch_id,
                "status": "unknown",
                "error_message": str(e)
            }