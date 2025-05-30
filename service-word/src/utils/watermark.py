import os
import tempfile
import logging
from typing import Optional, Dict, Any
from docx import Document
import io

try:
    WATERMARK_METHOD = None
    
    try:
        from docx_watermark import add_watermark
        WATERMARK_METHOD = "docx-watermark"
    except ImportError:
        pass
        
    if not WATERMARK_METHOD:
        try:
            import groupdocs.watermark as gw
            import groupdocs.watermark.watermarks as gwo
            WATERMARK_METHOD = "groupdocs"
        except ImportError:
            pass
            
    if not WATERMARK_METHOD:
        WATERMARK_METHOD = "direct"
        
except Exception:
    WATERMARK_METHOD = "direct"
    
logger = logging.getLogger(__name__)


class WatermarkHelper:
    """
    Lớp trợ giúp thêm watermark vào tài liệu Word.
    """
    
    @staticmethod
    def add_watermark(
        input_data: bytes, 
        text: str, 
        position: str = "center", 
        opacity: float = 0.5,
        font_name: str = "Times New Roman",
        font_size: int = 40,
        rotation: int = -45
    ) -> bytes:
        """
        Thêm watermark vào tài liệu Word.
        
        Args:
            input_data: Dữ liệu đầu vào của tài liệu Word
            text: Nội dung watermark
            position: Vị trí của watermark (center, top-left, top-right, bottom-left, bottom-right)
            opacity: Độ mờ của watermark (0.0 - 1.0)
            font_name: Tên font
            font_size: Kích thước font
            rotation: Góc xoay (độ)
            
        Returns:
            Dữ liệu tài liệu Word đã thêm watermark
        """
        logger.info(f"Thêm watermark '{text}' vào tài liệu bằng phương pháp: {WATERMARK_METHOD}")
        
        try:
            if WATERMARK_METHOD == "docx-watermark":
                return WatermarkHelper._add_watermark_with_docx_watermark(
                    input_data, text, position, opacity, font_name, font_size, rotation
                )
            elif WATERMARK_METHOD == "groupdocs":
                return WatermarkHelper._add_watermark_with_groupdocs(
                    input_data, text, position, opacity, font_name, font_size, rotation
                )
            else:
                return WatermarkHelper._add_watermark_direct(
                    input_data, text, position, opacity, font_name, font_size, rotation
                )
        except Exception as e:
            logger.error(f"Lỗi khi thêm watermark: {str(e)}")
            raise
            
    @staticmethod
    def _add_watermark_with_docx_watermark(
        input_data: bytes,
        text: str,
        position: str = "center",
        opacity: float = 0.5,
        font_name: str = "Times New Roman",
        font_size: int = 40,
        rotation: int = -45
    ) -> bytes:
        """
        Thêm watermark bằng thư viện python-docx-watermark.
        """
        from docx_watermark import add_watermark
        
        with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as temp_in:
            temp_in.write(input_data)
            temp_in_path = temp_in.name
        
        with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as temp_out:
            temp_out_path = temp_out.name
            
        try:
            add_watermark(
                input_docx=temp_in_path,
                output_docx=temp_out_path,
                text=text,
                opacity=opacity,
                size=font_size,
                rotation=rotation,
                italic=False
            )
            
            with open(temp_out_path, "rb") as f:
                output_data = f.read()
                
            return output_data
            
        finally:
            if os.path.exists(temp_in_path):
                os.unlink(temp_in_path)
            if os.path.exists(temp_out_path):
                os.unlink(temp_out_path)
                
    @staticmethod
    def _add_watermark_with_groupdocs(
        input_data: bytes,
        text: str,
        position: str = "center",
        opacity: float = 0.5,
        font_name: str = "Times New Roman",
        font_size: int = 40,
        rotation: int = -45
    ) -> bytes:
        """
        Thêm watermark bằng thư viện groupdocs-watermark.
        """
        import groupdocs.watermark as gw
        import groupdocs.watermark.watermarks as gwo
        
        with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as temp_in:
            temp_in.write(input_data)
            temp_in_path = temp_in.name
            
        with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as temp_out:
            temp_out_path = temp_out.name
            
        try:
            font = gwo.Font(font_name, font_size)
            
            with gw.Watermarker(temp_in_path) as watermarker:
                watermark = gwo.TextWatermark(text, font)
                
                if position == "center":
                    watermark.x = 0
                    watermark.y = 0
                    watermark.horizontal_alignment = gwo.WatermarkAlignment.CENTER
                    watermark.vertical_alignment = gwo.WatermarkAlignment.CENTER
                elif position == "top-left":
                    watermark.x = 10
                    watermark.y = 10
                    watermark.horizontal_alignment = gwo.WatermarkAlignment.LEFT
                    watermark.vertical_alignment = gwo.WatermarkAlignment.TOP
                elif position == "top-right":
                    watermark.x = -10
                    watermark.y = 10
                    watermark.horizontal_alignment = gwo.WatermarkAlignment.RIGHT
                    watermark.vertical_alignment = gwo.WatermarkAlignment.TOP
                elif position == "bottom-left":
                    watermark.x = 10
                    watermark.y = -10
                    watermark.horizontal_alignment = gwo.WatermarkAlignment.LEFT
                    watermark.vertical_alignment = gwo.WatermarkAlignment.BOTTOM
                elif position == "bottom-right":
                    watermark.x = -10
                    watermark.y = -10
                    watermark.horizontal_alignment = gwo.WatermarkAlignment.RIGHT
                    watermark.vertical_alignment = gwo.WatermarkAlignment.BOTTOM
                    
                watermark.rotate_angle = rotation
                watermark.opacity = opacity
                
                watermarker.add(watermark)
                watermarker.save(temp_out_path)

            with open(temp_out_path, "rb") as f:
                output_data = f.read()
                
            return output_data
            
        finally:
            if os.path.exists(temp_in_path):
                os.unlink(temp_in_path)
            if os.path.exists(temp_out_path):
                os.unlink(temp_out_path)
                
    @staticmethod
    def _add_watermark_direct(
        input_data: bytes,
        text: str,
        position: str = "center",
        opacity: float = 0.5,
        font_name: str = "Times New Roman",
        font_size: int = 40,
        rotation: int = -45
    ) -> bytes:
        """
        Thêm watermark bằng python-docx trực tiếp.
        Đây là phương pháp dự phòng, hiệu quả kém hơn.
        """
        from docx.shared import RGBColor, Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Inches
        
        doc = Document(io.BytesIO(input_data))
        
        for section in doc.sections:
            header = section.header
            paragraph = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            run = paragraph.add_run(text)
            font = run.font
            font.name = font_name
            font.size = Pt(font_size)
            
            gray_level = int(170 + (255 - 170) * (1 - opacity))
            font.color.rgb = RGBColor(gray_level, gray_level, gray_level)
        
        output_buffer = io.BytesIO()
        doc.save(output_buffer)
        output_buffer.seek(0)
        
        return output_buffer.read() 