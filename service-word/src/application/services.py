import os
import io
import tempfile
import asyncio
import uuid
import json
import pandas as pd
import zipfile
from typing import List, Dict, Any, Optional, Tuple
from datetime import datetime
import logging
import shutil
from fastapi import BackgroundTasks

from application.dto import CreateDocumentDTO, CreateTemplateDTO, TemplateDataDTO, WatermarkDTO, BatchProcessingDTO, InternshipReportModel, RewardReportModel, LaborContractModel
from domain.models import WordDocumentInfo as DocumentInfo, TemplateInfo, BatchProcessingInfo
from domain.exceptions import DocumentNotFoundException, TemplateNotFoundException, StorageException
from domain.exceptions import ConversionException, WatermarkException, TemplateApplicationException, InvalidDataFormatException
from infrastructure.repository import DocumentRepository, TemplateRepository, BatchProcessingRepository
from infrastructure.minio_client import MinioClient
from infrastructure.rabbitmq_client import RabbitMQClient
from core.config import settings
from utils import WordConverter
from utils.watermark import WatermarkHelper

from docx import Document as DocxDocument
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm

logger = logging.getLogger(__name__)

class DocumentService:
    """
    Service xử lý tài liệu Word.
    """

    def __init__(
            self,
            document_repository: DocumentRepository,
            minio_client: MinioClient,
            rabbitmq_client: RabbitMQClient
    ):
        """
        Khởi tạo service.

        Args:
            document_repository: Repository để làm việc với tài liệu
            minio_client: Client MinIO để lưu trữ tài liệu
            rabbitmq_client: Client RabbitMQ để gửi tin nhắn
        """
        self.document_repository = document_repository
        self.minio_client = minio_client
        self.rabbitmq_client = rabbitmq_client

    async def create_document(self, dto: CreateDocumentDTO, content: bytes) -> DocumentInfo:
        """
        Tạo tài liệu mới.

        Args:
            dto: DTO chứa thông tin tài liệu
            content: Nội dung tài liệu

        Returns:
            Thông tin tài liệu đã tạo
        """
        document_info = DocumentInfo(
            title=dto.title,
            description=dto.description,
            original_filename=dto.original_filename,
            file_size=len(content),
            file_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document" if dto.original_filename.endswith(
                ".docx") else "application/msword",
            storage_path="",  
            doc_metadata=dto.doc_metadata,
            user_id=dto.user_id
        )

        document_info = await self.document_repository.save(document_info, content)

        return document_info

    async def get_documents(self, skip: int = 0, limit: int = 10, search: Optional[str] = None, user_id: Optional[str] = None) -> Tuple[List[DocumentInfo], int]:
        """
        Lấy danh sách tài liệu.

        Args:
            skip: Số tài liệu bỏ qua
            limit: Số tài liệu tối đa trả về
            search: Từ khóa tìm kiếm
            user_id: ID của người dùng để lọc tài liệu

        Returns:
            Tuple chứa danh sách tài liệu và tổng số tài liệu
        """
        return await self.document_repository.list(skip, limit, search, user_id)

    async def get_document(self, document_id: str, user_id_check: Optional[str] = None) -> Tuple[DocumentInfo, bytes]:
        """
        Lấy thông tin và nội dung tài liệu.

        Args:
            document_id: ID của tài liệu
            user_id_check: ID người dùng để kiểm tra quyền sở hữu (optional)

        Returns:
            Tuple chứa thông tin và nội dung tài liệu
        """
        return await self.document_repository.get(document_id, user_id_check=user_id_check)

    async def delete_document(self, document_id: str, user_id_check: Optional[str] = None) -> None:
        """
        Xóa tài liệu.

        Args:
            document_id: ID của tài liệu
            user_id_check: ID người dùng để kiểm tra quyền sở hữu (optional)
        """
        await self.document_repository.delete(document_id, user_id_check=user_id_check)

    async def download_document(self, document_id: str, user_id_check: Optional[str] = None) -> Tuple[str, str, str]:
        """
        Tải xuống tài liệu.

        Args:
            document_id: ID của tài liệu
            user_id_check: ID người dùng để kiểm tra quyền sở hữu (optional)

        Returns:
            Tuple chứa (file_path, media_type, filename)
        """
        try:
            document_info, content = await self.document_repository.get(document_id, user_id_check=user_id_check)
            
            # Tạo file tạm thời để trả về
            temp_dir = settings.TEMP_DIR
            os.makedirs(temp_dir, exist_ok=True)
            
            temp_filename = f"{uuid.uuid4()}_{document_info.original_filename}"
            temp_path = os.path.join(temp_dir, temp_filename)
            
            with open(temp_path, "wb") as f:
                f.write(content)
            
            return temp_path, document_info.file_type, document_info.original_filename
            
        except Exception as e:
            logger.error(f"Lỗi khi tải tài liệu {document_id}: {e}", exc_info=True)
            raise StorageException(f"Không thể tải tài liệu: {str(e)}")

    async def convert_to_pdf(self, content: bytes, original_filename: str, user_id: str) -> Dict[str, Any]:
        """
        Chuyển đổi tài liệu Word sang PDF.

        Args:
            content: Nội dung tài liệu Word
            original_filename: Tên file gốc
            user_id: ID của người dùng

        Returns:
            Dict chứa thông tin tài liệu PDF
        """
        try:
            with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as temp_word_file_obj:
                temp_word_file_obj.write(content)
                temp_word_path = temp_word_file_obj.name
            
            pdf_basename = os.path.splitext(original_filename)[0] + ".pdf"
            os.makedirs(settings.TEMP_DIR, exist_ok=True)
            temp_pdf_path = os.path.join(settings.TEMP_DIR, f"{uuid.uuid4()}_{pdf_basename}")

            try:
                WordConverter.convert_to_pdf(
                    input_path=temp_word_path,
                    output_path=temp_pdf_path,
                    method="libreoffice"  
                )

                with open(temp_pdf_path, "rb") as f:
                    pdf_content = f.read()
            finally:
                if os.path.exists(temp_word_path):
                    os.unlink(temp_word_path)
                if os.path.exists(temp_pdf_path) and not 'pdf_content' in locals():
                    os.unlink(temp_pdf_path)

            document_info_pdf = DocumentInfo(
                title=os.path.splitext(original_filename)[0] + " (PDF)",
                description=f"PDF được chuyển đổi từ {original_filename}",
                original_filename=pdf_basename,
                file_size=len(pdf_content),
                file_type="application/pdf",
                storage_path="",  
                doc_metadata={"converted_from": original_filename, "original_word_filename": original_filename},
                user_id=user_id
            )
            saved_pdf_info = await self.document_repository.save(document_info_pdf, pdf_content)
            
            if os.path.exists(temp_pdf_path):
                os.unlink(temp_pdf_path)

            return {
                "id": saved_pdf_info.id,
                "filename": saved_pdf_info.original_filename,
                "file_size": saved_pdf_info.file_size
            }
        except Exception as e:
            logger.error(f"Lỗi khi chuyển đổi Word sang PDF cho file {original_filename}, user {user_id}: {e}", exc_info=True)
            if 'temp_word_path' in locals() and os.path.exists(temp_word_path):
                os.unlink(temp_word_path)
            if 'temp_pdf_path' in locals() and os.path.exists(temp_pdf_path):
                os.unlink(temp_pdf_path)
            raise ConversionException(f"Lỗi khi chuyển đổi sang PDF: {str(e)}")

    async def add_watermark(self, content: bytes, original_filename: str, dto: WatermarkDTO, user_id: str) -> Dict[str, Any]:
        """
        Thêm watermark vào tài liệu Word.

        Args:
            content: Nội dung tài liệu Word
            original_filename: Tên file gốc
            dto: DTO chứa thông tin watermark
            user_id: ID của người dùng

        Returns:
            Dict chứa thông tin tài liệu đã thêm watermark
        """
        try:
            output_content = WatermarkHelper.add_watermark(
                input_data=content,
                text=dto.text,
                position=dto.position,
                opacity=dto.opacity,
                font_name=dto.font_name,
                font_size=dto.font_size,
                rotation=dto.rotation
            )

            parts = os.path.splitext(original_filename)
            watermarked_filename = f"{parts[0]}_watermarked{parts[1]}"
    
            document_info_wm = DocumentInfo(
                title=f"{parts[0]} (Watermarked)",
                description=f"Tài liệu đã thêm watermark '{dto.text}'",
                original_filename=watermarked_filename,
                file_size=len(output_content),
                file_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document" if watermarked_filename.endswith(".docx") else "application/msword",
                storage_path="",
                doc_metadata={
                    "watermark": dto.text,
                    "watermark_position": dto.position,
                    "original_file": original_filename
                },
                user_id=user_id
            )
            
            saved_wm_info = await self.document_repository.save(document_info_wm, output_content)
            
            return {
                "id": saved_wm_info.id,
                "filename": saved_wm_info.original_filename,
                "file_size": saved_wm_info.file_size
            }
        except Exception as e:
            logger.error(f"Lỗi khi thêm watermark vào file {original_filename}, user {user_id}: {str(e)}", exc_info=True)
            raise

    async def process_document_async(self, document_id: str, user_id: Optional[str] = None) -> None:
        """
        Xử lý tài liệu bất đồng bộ.

        Args:
            document_id: ID của tài liệu
        """
        try:
            document_info, _ = await self.document_repository.get(document_id, user_id_check=user_id)

            await self.rabbitmq_client.publish_convert_to_pdf_task(document_id)
        except Exception as e:
            print(f"Lỗi khi xử lý tài liệu {document_id}: {str(e)}")


class TemplateService:
    """
    Service xử lý mẫu tài liệu Word.
    """

    def __init__(
            self,
            template_repository: TemplateRepository,
            minio_client: MinioClient,
            rabbitmq_client: RabbitMQClient,
            batch_processing_repository: Optional[BatchProcessingRepository] = None
    ):
        """
        Khởi tạo service.

        Args:
            template_repository: Repository để làm việc với mẫu tài liệu
            minio_client: Client MinIO để lưu trữ mẫu tài liệu
            rabbitmq_client: Client RabbitMQ để gửi tin nhắn
        """
        self.template_repository = template_repository
        self.minio_client = minio_client
        self.rabbitmq_client = rabbitmq_client
        self.document_repository = None
        self.batch_processing_repository = batch_processing_repository or BatchProcessingRepository()

    def set_document_repository(self, document_repository: DocumentRepository):
        """Cho phép inject DocumentRepository sau khi khởi tạo, tránh circular dependency."""
        self.document_repository = document_repository

    async def create_template(self, dto: CreateTemplateDTO, content: bytes) -> TemplateInfo:
        """
        Tạo mẫu tài liệu mới.

        Args:
            dto: DTO chứa thông tin mẫu tài liệu
            content: Nội dung mẫu tài liệu

        Returns:
            Thông tin mẫu tài liệu đã tạo
        """
        template_info = TemplateInfo(
            template_id=dto.template_id or str(uuid.uuid4()),
            name=dto.name,
            description=dto.description,
            category=dto.category,
            tags=dto.tags or [],
            original_filename=dto.original_filename,
            file_size=len(content),
            doc_metadata=dto.doc_metadata
        )
        return await self.template_repository.save(template_info, content)

    async def get_templates(self, category: Optional[str] = None, skip: int = 0, limit: int = 10) -> Tuple[List[TemplateInfo], int]:
        """
        Lấy danh sách mẫu tài liệu.

        Args:
            category: Danh mục để lọc
            skip: Số mẫu tài liệu bỏ qua
            limit: Số mẫu tài liệu tối đa trả về

        Returns:
            Tuple chứa danh sách mẫu tài liệu và tổng số mẫu tài liệu
        """
        return await self.template_repository.list(category, skip, limit)

    async def get_template(self, template_id: str) -> Tuple[TemplateInfo, bytes]:
        """
        Lấy thông tin và nội dung mẫu tài liệu.

        Args:
            template_id: ID của mẫu tài liệu

        Returns:
            Tuple chứa thông tin và nội dung mẫu tài liệu
        """
        return await self.template_repository.get(template_id)

    async def delete_template(self, template_id: str) -> None:
        """
        Xóa mẫu tài liệu.

        Args:
            template_id: ID của mẫu tài liệu
        """
        await self.template_repository.delete(template_id)

    async def _save_applied_template_as_document(
        self, 
        content_bytes: bytes, 
        original_template_filename: str, 
        output_format: str, 
        user_id: str, 
        template_id: str,
        filled_data: Dict[str, Any]
    ) -> Dict[str, Any]:
        """Lưu tài liệu được tạo từ template vào DocumentRepository."""
        if not self.document_repository:
            raise Exception("DocumentRepository is not set in TemplateService. Cannot save applied template.")

        base_name = os.path.splitext(original_template_filename)[0]
        output_filename = f"{base_name}_applied.{output_format}"
        
        file_type = "application/pdf" if output_format == "pdf" else (
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document" if output_format == "docx" else "application/msword")

        doc_info_to_save = DocumentInfo(
            title=f"{base_name} (from template {template_id})",
            description=f"Created from template {template_id}",
            original_filename=output_filename,
            file_size=len(content_bytes),
            file_type=file_type,
            document_category="word",
            user_id=user_id,
            doc_metadata={
                "template_id": template_id,
                "template_source_filename": original_template_filename,
                "filled_data_preview": {k: str(v)[:50] + '...' if isinstance(v, str) and len(v) > 50 else v for k, v in filled_data.items()}
            }
        )
        saved_doc = await self.document_repository.save(doc_info_to_save, content_bytes)
        return {
            "id": saved_doc.id,
            "filename": saved_doc.original_filename,
            "file_size": saved_doc.file_size
        }

    async def apply_template(self, dto: TemplateDataDTO) -> Dict[str, Any]:
        template_info, template_content = await self.template_repository.get(dto.template_id)
        if not template_info:
            raise TemplateNotFoundException(dto.template_id)

        filled_content_bytes = template_content

        if template_info.original_filename.endswith(('.doc', '.docx')) and dto.output_format == 'pdf':
            try:
                with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as temp_filled_doc:
                    temp_filled_doc.write(filled_content_bytes)
                    temp_filled_doc_path = temp_filled_doc.name
                
                pdf_output_path = os.path.splitext(temp_filled_doc_path)[0] + ".pdf"
                WordConverter.convert_to_pdf(temp_filled_doc_path, pdf_output_path, method="libreoffice")
                
                with open(pdf_output_path, "rb") as f:
                    final_output_bytes = f.read()
                
                os.unlink(temp_filled_doc_path)
                os.unlink(pdf_output_path)
            except Exception as e:
                logger.error(f"Error converting filled template {dto.template_id} to PDF for user {dto.user_id}: {e}", exc_info=True)
                raise TemplateApplicationException(f"Error converting filled template to PDF: {e}")
        else:
            final_output_bytes = filled_content_bytes

        return await self._save_applied_template_as_document(
            final_output_bytes, 
            template_info.original_filename, 
            dto.output_format, 
            dto.user_id,
            dto.template_id,
            dto.data
        )

    async def _create_document_from_data(self, data: Dict[str, Any], template_filename: str, output_filename_base: str, user_id: str) -> Tuple[str, bytes]:
        """
        Hàm helper để tạo một tài liệu Word từ dữ liệu (ví dụ cho báo cáo, hợp đồng).
        template_filename: tên file template trong thư mục templates/core (ví dụ: 'internship_report_template.docx')
        output_filename_base: tên file output (không có extension, ví dụ: 'internship_report_john_doe')
        data: dictionary chứa dữ liệu để điền vào template.
        user_id: id của người dùng.
        Trả về: (tên file output thực tế, nội dung file bytes)
        """
        logger.info(f"Tạo tài liệu {output_filename_base} cho user {user_id} từ template {template_filename}")
        doc = DocxDocument()
        doc.add_heading(f'Placeholder cho {output_filename_base}', 0)
        doc.add_paragraph(f'Dữ liệu: {json.dumps(data, ensure_ascii=False, indent=2)}')
        output_io = io.BytesIO()
        doc.save(output_io)
        output_bytes = output_io.getvalue()
        generated_filename = f"{output_filename_base}.docx"
        return generated_filename, output_bytes

    async def create_internship_report(self, data: InternshipReportModel, user_id: str) -> Dict[str, Any]:
        output_filename_base = f"internship_report_{data.intern_name.replace(' ', '_')}_{datetime.now().strftime('%Y%m%d')}"
        generated_filename, output_bytes = await self._create_document_from_data(
            data.dict(), 
            "internship_report_template.docx",
            output_filename_base,
            user_id
        )
        return await self._save_applied_template_as_document(
            output_bytes, 
            "internship_report_template.docx", 
            "docx", 
            user_id,
            "internship-report-form",
            data.dict()
        )

    async def create_reward_report(self, data: RewardReportModel, user_id: str) -> Dict[str, Any]:
        output_filename_base = f"reward_report_{data.recipient.replace(' ', '_')}_{datetime.now().strftime('%Y%m%d')}"
        generated_filename, output_bytes = await self._create_document_from_data(
            data.dict(),
            "reward_report_template.docx",
            output_filename_base,
            user_id
        )
        return await self._save_applied_template_as_document(
            output_bytes,
            "reward_report_template.docx",
            "docx",
            user_id,
            "reward-report-form",
            data.dict()
        )

    async def create_labor_contract(self, data: LaborContractModel, user_id: str) -> Dict[str, Any]:
        output_filename_base = f"labor_contract_{data.employee_name.replace(' ', '_')}_{datetime.now().strftime('%Y%m%d')}"
        generated_filename, output_bytes = await self._create_document_from_data(
            data.dict(),
            "labor_contract_template.docx",
            output_filename_base,
            user_id
        )
        return await self._save_applied_template_as_document(
            output_bytes,
            "labor_contract_template.docx",
            "docx",
            user_id,
            "labor-contract-form",
            data.dict()
        )

    async def process_batch_async(self, 
                                  task_id: str, 
                                  template_id_or_name: str,
                                  data_list: List[Dict[str, Any]],
                                  original_data_filename: str,
                                  output_format: str, 
                                  user_id: str) -> None:
        """
        Xử lý tạo tài liệu hàng loạt một cách bất đồng bộ.
        Lưu từng tài liệu được tạo và cập nhật trạng thái của batch processing.
        """
        logger.info(f"Bắt đầu process_batch_async cho task_id: {task_id}, user_id: {user_id}")
        batch_info = BatchProcessingInfo(
            task_id=task_id,
            user_id=user_id,
            template_id=template_id_or_name,
            status="PROCESSING",
            total_files=len(data_list),
            processed_files=0,
            output_format=output_format,
            original_data_filename=original_data_filename,
            generated_documents=[]
        )
        if self.batch_processing_repository:
            await self.batch_processing_repository.save(batch_info)

        processed_docs_info = []
        temp_dir_for_batch = os.path.join(settings.TEMP_DIR, task_id)
        os.makedirs(temp_dir_for_batch, exist_ok=True)
        generated_files_paths = []

        try:
            for i, item_data in enumerate(data_list):
                try:
                    output_filename_base = f"{os.path.splitext(original_data_filename)[0]}_item_{i+1}"
                    _temp_doc = DocxDocument()
                    _temp_doc.add_heading(f'Placeholder cho {output_filename_base} (item {i+1})', 0)
                    _temp_doc.add_paragraph(f'Dữ liệu: {json.dumps(item_data, ensure_ascii=False, indent=2)}')
                    temp_output_io = io.BytesIO()
                    _temp_doc.save(temp_output_io)
                    item_output_bytes = temp_output_io.getvalue()
                    actual_output_filename_base = output_filename_base

                    if output_format == 'pdf':
                        with tempfile.NamedTemporaryFile(delete=False, suffix=".docx", dir=temp_dir_for_batch) as temp_filled_item_doc:
                            temp_filled_item_doc.write(item_output_bytes)
                            temp_filled_item_doc_path = temp_filled_item_doc.name
                        
                        pdf_item_output_path = os.path.splitext(temp_filled_item_doc_path)[0] + ".pdf"
                        WordConverter.convert_to_pdf(temp_filled_item_doc_path, pdf_item_output_path, method="libreoffice")
                        
                        with open(pdf_item_output_path, "rb") as f:
                            final_item_output_bytes = f.read()
                        
                        os.unlink(temp_filled_item_doc_path)
                        actual_output_filename_with_ext = f"{actual_output_filename_base}.pdf"
                        path_to_save_for_zip = pdf_item_output_path
                    else:
                        final_item_output_bytes = item_output_bytes
                        actual_output_filename_with_ext = f"{actual_output_filename_base}.docx"
                        temp_docx_path = os.path.join(temp_dir_for_batch, actual_output_filename_with_ext)
                        with open(temp_docx_path, "wb") as f_temp_docx:
                            f_temp_docx.write(final_item_output_bytes)
                        path_to_save_for_zip = temp_docx_path

                    generated_files_paths.append(path_to_save_for_zip)
                    
                    batch_info.processed_files += 1

                except Exception as e_item:
                    logger.error(f"Lỗi khi xử lý item {i+1} cho batch {task_id}: {e_item}", exc_info=True)
                    batch_info.errors.append(f"Item {i+1}: {str(e_item)}")
                
                if self.batch_processing_repository:
                    batch_info.status = "PROCESSING" if batch_info.processed_files < batch_info.total_files else "COMPLETED"
                    await self.batch_processing_repository.update(batch_info)

            if output_format == "zip" or (output_format == "pdf" and len(data_list) > 1) or (output_format == "docx" and len(data_list) > 1):
                zip_filename_base = f"{os.path.splitext(original_data_filename)[0]}_batch_{task_id}"
                zip_filename = f"{zip_filename_base}.zip"
                zip_output_path = os.path.join(settings.TEMP_DIR, zip_filename)

                with zipfile.ZipFile(zip_output_path, 'w', zipfile.ZIP_DEFLATED) as zipf:
                    for file_path_in_batch_dir in generated_files_paths:
                        zipf.write(file_path_in_batch_dir, os.path.basename(file_path_in_batch_dir))
                
                with open(zip_output_path, "rb") as f_zip:
                    zip_content_bytes = f_zip.read()
                
                zip_doc_info = DocumentInfo(
                    title=f"Batch Processed: {original_data_filename} (Task {task_id})",
                    description=f"Archive of {len(data_list)} documents generated from template {template_id_or_name} and data {original_data_filename}. Task ID: {task_id}",
                    original_filename=zip_filename,
                    file_size=len(zip_content_bytes),
                    file_type="application/zip",
                    document_category="archive",
                    user_id=user_id,
                    doc_metadata={
                        "batch_task_id": task_id,
                        "template_id": template_id_or_name,
                        "source_data_file": original_data_filename,
                        "num_items": len(data_list),
                        "output_format_items": output_format if output_format != "zip" else "docx"
                    }
                )
                if not self.document_repository:
                    raise Exception("DocumentRepository not set for TemplateService, cannot save batch ZIP.")
                saved_zip_doc = await self.document_repository.save(zip_doc_info, zip_content_bytes)
                batch_info.generated_documents.append({"document_id": saved_zip_doc.id, "filename": saved_zip_doc.original_filename})
                os.unlink(zip_output_path)
            
            elif len(generated_files_paths) == 1 and output_format != "zip":
                single_file_path = generated_files_paths[0]
                with open(single_file_path, "rb") as f_single:
                    single_file_bytes = f_single.read()
                
                single_doc_info = DocumentInfo(
                     title=f"Generated: {os.path.basename(single_file_path)} (Task {task_id})",
                     description=f"Document generated from template {template_id_or_name} and data {original_data_filename}. Task ID: {task_id}",
                     original_filename=os.path.basename(single_file_path),
                     file_size=len(single_file_bytes),
                     file_type="application/pdf" if output_format == "pdf" else "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                     document_category="word",
                     user_id=user_id,
                     doc_metadata={
                        "batch_task_id": task_id,
                        "template_id": template_id_or_name,
                        "source_data_file": original_data_filename
                     }
                )
                if not self.document_repository:
                    raise Exception("DocumentRepository not set for TemplateService, cannot save single batch item.")
                saved_single_doc = await self.document_repository.save(single_doc_info, single_file_bytes)
                batch_info.generated_documents.append({"document_id": saved_single_doc.id, "filename": saved_single_doc.original_filename})

            batch_info.status = "COMPLETED" if not batch_info.errors else "COMPLETED_WITH_ERRORS"
        except Exception as e_batch:
            logger.error(f"Lỗi nghiêm trọng trong process_batch_async cho task {task_id}: {e_batch}", exc_info=True)
            batch_info.status = "FAILED"
            batch_info.errors.append(f"Batch processing failed: {str(e_batch)}")
        finally:
            if self.batch_processing_repository:
                await self.batch_processing_repository.update(batch_info)
            if os.path.exists(temp_dir_for_batch):
                shutil.rmtree(temp_dir_for_batch, ignore_errors=True)
            logger.info(f"Kết thúc process_batch_async cho task_id: {task_id}, status: {batch_info.status}")

    def _parse_data_file(self, file_content: bytes, original_filename: str) -> List[Dict[str, Any]]:
        """Parse CSV or Excel file content into a list of dictionaries."""
        if original_filename.endswith(".csv"):
            try:
                data_io = io.StringIO(file_content.decode('utf-8'))
                df = pd.read_csv(data_io)
            except UnicodeDecodeError:
                data_io = io.BytesIO(file_content)
                df = pd.read_csv(data_io)
            except Exception as e:
                raise InvalidDataFormatException(f"Lỗi đọc file CSV: {original_filename}. {str(e)}")

        elif original_filename.endswith(('.xlsx', '.xls')):
            try:
                data_io = io.BytesIO(file_content)
                df = pd.read_excel(data_io)
            except Exception as e:
                raise InvalidDataFormatException(f"Lỗi đọc file Excel: {original_filename}. {str(e)}")
        else:
            raise InvalidDataFormatException(f"Định dạng file dữ liệu không được hỗ trợ: {original_filename}. Chỉ chấp nhận CSV hoặc Excel.")
        
        df = df.fillna('')
        return df.to_dict(orient='records')

    async def create_batch_documents_from_file(self, 
                                               template_id: str, 
                                               file_content: bytes, 
                                               original_filename: str,
                                               output_format: str, 
                                               user_id: str,
                                               background_tasks: BackgroundTasks) -> str:
        """
        Tạo tài liệu hàng loạt từ một template và một file dữ liệu (CSV/Excel).
        Sử dụng background task để xử lý.
        """
        try:
            data_list = self._parse_data_file(file_content, original_filename)
        except InvalidDataFormatException as e:
            logger.error(f"Lỗi parse file dữ liệu {original_filename} cho batch: {e}")
            raise

        if not data_list:
            raise InvalidDataFormatException(f"Không có dữ liệu trong file: {original_filename}")

        task_id = str(uuid.uuid4())
        
        background_tasks.add_task(
            self.process_batch_async,
            task_id=task_id,
            template_id_or_name=template_id,
            data_list=data_list,
            original_data_filename=original_filename,
            output_format=output_format,
            user_id=user_id
        )
        logger.info(f"Đã tạo task xử lý batch (create_batch_documents_from_file) với ID: {task_id} cho user {user_id}")
        return task_id

    async def generate_invitations_from_file(self,
                                             file_content: bytes,
                                             original_filename: str,
                                             output_format: str,
                                             user_id: str,
                                             background_tasks: BackgroundTasks) -> str:
        """
        Tạo lời mời hàng loạt từ template 'invitation' và file dữ liệu.
        """
        INVITATION_TEMPLATE_ID = "invitation_template.docx"
        logger.info(f"Bắt đầu generate_invitations_from_file cho user {user_id} từ file {original_filename}")
        
        try:
            data_list = self._parse_data_file(file_content, original_filename)
        except InvalidDataFormatException as e:
            logger.error(f"Lỗi parse file dữ liệu {original_filename} cho invitations: {e}")
            raise 

        if not data_list:
            raise InvalidDataFormatException(f"Không có dữ liệu trong file: {original_filename}")

        task_id = str(uuid.uuid4())
        
        background_tasks.add_task(
            self.process_batch_async,
            task_id=task_id,
            template_id_or_name=INVITATION_TEMPLATE_ID, 
            data_list=data_list,
            original_data_filename=original_filename,
            output_format=output_format,
            user_id=user_id
        )
        logger.info(f"Đã tạo task xử lý batch (generate_invitations_from_file) với ID: {task_id} cho user {user_id}")
        return task_id

    async def get_batch_status(self, batch_id: str, user_id: Optional[str] = None) -> Optional[BatchProcessingInfo]:
        """
        Lấy trạng thái của một tác vụ xử lý batch.
        Kiểm tra user_id nếu BatchProcessingRepository hỗ trợ.
        """
        if not self.batch_processing_repository:
            logger.warning("BatchProcessingRepository is not available in TemplateService.")
            return None
        return await self.batch_processing_repository.get(batch_id)